# -*- coding: utf-8 -*-
"""Получение информации по IP, доменам, BIN и номерам телефонов."""

import html
import os
import re
import logging
from typing import Optional, Any

# Экранирование для HTML (Telegram parse_mode)
def _h(val: Any) -> str:
    """Экранирует строку для безопасного вывода в HTML."""
    if val is None:
        return "—"
    return html.escape(str(val).strip()) if str(val).strip() else "—"

# Разделитель (на всю ширину сообщения в Telegram)
DIV = "─" * 30
E_IP, E_DOMAIN, E_BIN, E_PHONE, E_WALLET = "📡", "📌", "💳", "📱", "₿"
E_COUNTRY, E_CITY, E_ORG, E_DATE, E_LINK = "🌍", "📍", "🏢", "📅", "🌐"
E_SECURITY, E_SCHEME, E_TYPE = "🔒", "🏦", "📋"
E_STAR, E_WARN, E_TIME = "⭐", "⚠️", "🕐"

import requests

try:
    import dns.resolver
    DNS_AVAILABLE = True
except ImportError:
    DNS_AVAILABLE = False

logger = logging.getLogger(__name__)

# Реквизиты для справки (заполните для подстановки, иначе будут подчёркивания)
SPRAVKA_CASE = ""  # № уголовного дела, например "123456/2025"
SPRAVKA_POSITION = ""  # например "Следователь ОМВД России"
SPRAVKA_RANK = ""  # например "старший лейтенант юстиции"
SPRAVKA_NAME = ""  # например "И.И.Иванов"

API_URL = "https://api.ipapi.is/"
WHOISJSON_URL = "https://whoisjson.com/api/v1/whois"
RDAP_DOMAIN_URL = "https://rdap.org/domain"
VIEWDNS_WHOIS_URL = "https://viewdns.info/whois/"
HANDYAPI_BIN_URL = "https://data.handyapi.com/bin"
# Резервный BIN без ключа: lookup.binlist.net (лимит ~5 запросов/час, см. binlist.net).
BINLIST_LOOKUP_URL = "https://lookup.binlist.net"
# Ключи внешних API берём из переменных окружения, чтобы не хранить секреты в репозитории.
HANDYAPI_KEY = os.environ.get("HANDYAPI_KEY", "")
WHOISJSON_KEY = os.environ.get("WHOISJSON_KEY", "")
NUMVERIFY_URL = "https://apilayer.net/api/validate"
NUMVERIFY_KEY = os.environ.get("NUMVERIFY_KEY", "")
# Криптокошелёк: ETH (Etherscan), BTC (BlockCypher), TRON (TronGrid)
ETHERSCAN_API_URL = "https://api.etherscan.io/api"
ETHERSCAN_KEY = os.environ.get("ETHERSCAN_KEY", "")
BLOCKCYPHER_URL = "https://api.blockcypher.com/v1"
TRONGRID_URL = "https://api.trongrid.io"

# Метки известных адресов (биржи, миксеры) — lowercase
WALLET_LABELS: dict[str, str] = {
    "0x28c6c06298d514db089934071355e5743bf21d60": "Binance Hot Wallet",
    "0xf977814e90da44bfa03b6295a0616a897441acec": "Binance",
    "0x3f5ce5fbfe3e9af3971dd833d26ba9b5c936f0be": "Binance",
    "0xd551234ae421e3bcba99a0da6d736074f22192ff": "Binance",
    "0x722122df12d4e14e13ac3b6895a86e84145b6967": "Tornado Cash",
    "0xd90e2f925da726b50c4ed8d0fb90ad053324f31b": "Tornado Cash",
    "0x910cbd523d972eb0a6f4cae4618ad62622b39dbf": "Tornado Cash",
    "tr7nhqjekqxgtci8q8zy4pl8otszgjlj6t": "USDT (TRON)",
}


def _validate_ip(ip: str) -> bool:
    """Проверка формата IPv4 или IPv6."""
    ip = ip.strip()
    ipv4 = re.match(
        r"^(?:(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9])\.){3}"
        r"(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9])$",
        ip,
    )
    if ipv4:
        return True
    if ":" in ip and len(ip) <= 45:
        return True
    return False


def _looks_like_ipv4(candidate: str) -> bool:
    """Строка похожа на IPv4 (четыре числа через точку), даже если октеты невалидны (например 352)."""
    candidate = candidate.strip()
    if not re.match(r"^\d+\.\d+\.\d+\.\d+$", candidate):
        return False
    return True


def _yes_no(value: bool) -> str:
    return "Да" if value else "Нет"


def _format_ip_info(data: dict[str, Any], ip: str) -> str:
    """Форматирует ответ API в HTML."""
    loc = data.get("location") or {}
    asn_data = data.get("asn") or {}
    country = loc.get("country", "—")
    city = loc.get("city", "—")
    org = asn_data.get("org", "—")
    asn_num = asn_data.get("asn", "")
    provider = f"{org} ({asn_num})" if asn_num else org
    ip_val = data.get("ip", ip)

    lines = [
        f"{E_IP} <b>Информация об IP</b>: <code>{_h(ip_val)}</code>",
        DIV,
        f"{E_COUNTRY} <b>Страна:</b> {_h(country)}",
        f"{E_CITY} <b>Город:</b> {_h(city)}",
        f"{E_ORG} <b>Провайдер (ASN):</b> {_h(provider)}",
        "",
        f"{E_SECURITY} <b>Безопасность:</b>",
        f"  • Дата-центр/Хостинг: {_yes_no(data.get('is_datacenter', False))}",
        f"  • VPN: {_yes_no(data.get('is_vpn', False))}",
        f"  • Прокси: {_yes_no(data.get('is_proxy', False))}",
        f"  • TOR: {_yes_no(data.get('is_tor', False))}",
    ]
    return "\n".join(lines)


def get_ip_info(ip: str) -> str:
    """
    Запрашивает информацию по IP через ipapi.is и возвращает
    отформатированную строку заданного вида.
    """
    ip = ip.strip()
    if not _validate_ip(ip):
        if _looks_like_ipv4(ip):
            return "Неверный IP-адрес: один из октетов вне диапазона 0–255. Проверьте адрес (например 46.135.35.214)."
        return "Неверный формат IP. Укажите IPv4 или IPv6 адрес."

    try:
        r = requests.get(API_URL, params={"q": ip}, timeout=15)
        r.raise_for_status()
        data = r.json()
    except requests.RequestException as e:
        logger.warning("ipapi.is request failed: %s", e)
        return "Не удалось получить данные по этому IP. Попробуйте позже."
    except (ValueError, KeyError) as e:
        logger.warning("ipapi.is parse error: %s", e)
        return "Ошибка при разборе ответа API. Попробуйте позже."

    return _format_ip_info(data, ip)


def _fetch_ip_data_raw(ip: str) -> Optional[dict[str, Any]]:
    """Возвращает сырые данные ipapi.is для IP или None."""
    ip = ip.strip()
    if not _validate_ip(ip):
        return None
    try:
        r = requests.get(API_URL, params={"q": ip}, timeout=15)
        r.raise_for_status()
        return r.json()
    except Exception:
        return None


def _validate_domain(domain: str) -> bool:
    """Простая проверка формата доменного имени."""
    domain = domain.strip().lower()
    # Очень упрощённая проверка: имя.зона, без поддержи IDN
    return bool(
        re.match(
            r"^(?:[a-z0-9-]{1,63}\.)+[a-z]{2,63}$",
            domain,
        )
    )


def _validate_bin(bin_str: str) -> bool:
    """Проверка формата BIN: 6–8 цифр."""
    bin_str = bin_str.strip()
    return bool(re.match(r"^\d{6,8}$", bin_str))


def _validate_eth_address(addr: str) -> bool:
    """Ethereum: 0x + 40 hex."""
    addr = addr.strip()
    return bool(re.match(r"^0x[a-fA-F0-9]{40}$", addr))


def _validate_btc_address(addr: str) -> bool:
    """Bitcoin: 1..., 3..., bc1..."""
    addr = addr.strip()
    if re.match(r"^[13][a-km-zA-HJ-NP-Z1-9]{25,34}$", addr):
        return True
    if re.match(r"^bc1[a-zA-HJ-NP-Z0-9]{25,62}$", addr):
        return True
    return False


def _validate_tron_address(addr: str) -> bool:
    """TRON: T + base58, 34 символа."""
    addr = addr.strip()
    return bool(re.match(r"^T[a-km-zA-HJ-NP-Z1-9]{33}$", addr))


def _validate_wallet(addr: str) -> bool:
    """Проверка формата криптокошелька (ETH/BTC/TRON)."""
    return _validate_eth_address(addr) or _validate_btc_address(addr) or _validate_tron_address(addr)


def _detect_wallet_chains(addr: str) -> list[str]:
    """Возвращает список сетей, подходящих под адрес: eth, btc, tron."""
    chains = []
    if _validate_eth_address(addr):
        chains.append("eth")
    if _validate_btc_address(addr):
        chains.append("btc")
    if _validate_tron_address(addr):
        chains.append("tron")
    return chains


def _validate_phone(phone: str) -> bool:
    """Проверка формата номера телефона: 10–15 цифр (после удаления лишних символов)."""
    digits = re.sub(r"\D", "", phone.strip())
    return 10 <= len(digits) <= 15


def _normalize_phone(phone: str) -> str:
    """Нормализует номер: оставляет только цифры, добавляет + для международного формата."""
    digits = re.sub(r"\D", "", phone.strip())
    if not digits:
        return ""
    # 8xxxxxxxxxx (РФ) -> +7xxxxxxxxxx
    if digits.startswith("8") and len(digits) == 11:
        digits = "7" + digits[1:]
    return "+" + digits


def detect_lookup_type(text: str) -> tuple[str, str]:
    """
    Определяет тип запроса: IP, домен, BIN, телефон, wallet.
    Возвращает ("ip"|"domain"|"bin"|"phone"|"wallet", значение) или ("", "").
    """
    text = text.strip()
    parts = text.split()
    if len(parts) != 1:
        return "", ""
    candidate = parts[0].strip()
    if _validate_ip(candidate):
        return "ip", candidate
    # Похоже на IPv4 с опечаткой (например 46.135.352.114) — не считать телефоном
    if _looks_like_ipv4(candidate):
        return "ip", candidate
    if _validate_wallet(candidate):
        return "wallet", candidate
    if _validate_domain(candidate):
        return "domain", candidate
    if _validate_bin(candidate):
        return "bin", candidate
    if _validate_phone(candidate):
        return "phone", _normalize_phone(candidate)
    return "", ""


def _str_or_first(val: Any) -> str:
    """Преобразует значение в строку; если список — берёт первый элемент."""
    if val is None:
        return "—"
    if isinstance(val, list):
        return str(val[0]).strip() if val else "—"
    return str(val).strip() or "—"


def _format_date(val: Any) -> str:
    """Форматирует дату (datetime или строка) для вывода."""
    if val is None:
        return "—"
    if isinstance(val, list) and val:
        val = val[0]
    if hasattr(val, "strftime"):
        return val.strftime("%Y-%m-%d %H:%M")
    return str(val).strip() or "—"


def _get_nested(data: dict, *keys: str) -> Any:
    """Возвращает первое найденное значение по ключам (плоский или вложенный)."""
    for key in keys:
        if key in data and data[key] is not None:
            return data[key]
    # вложенные: registrant.country и т.п.
    for key in keys:
        if "." in key:
            part, rest = key.split(".", 1)
            if part in data and isinstance(data[part], dict):
                return _get_nested(data[part], rest)
    return None


def _extract_registrar_country(data: Any) -> Optional[str]:
    """
    Извлекает страну регистратора из ответа RDAP/WHOIS.
    RDAP: entities с role registrar -> vcardArray -> adr (последний элемент = страна).
    """
    if data is None:
        return None
    if isinstance(data, dict):
        # Прямые ключи
        direct = _get_nested(
            data,
            "registrarCountry",
            "registrar_country",
            "registrarCountryName",
            "registrar_country_name",
        )
        if direct is not None:
            s = _str_or_first(direct)
            if s and s != "—":
                return s
        # RDAP entities
        entities = data.get("entities") or data.get("registrar") or []
        if isinstance(entities, dict):
            entities = [entities]
        for entity in entities:
            if not isinstance(entity, dict):
                continue
            roles = entity.get("roles") or entity.get("role") or []
            if isinstance(roles, str):
                roles = [roles]
            if "registrar" not in roles:
                continue
            vcard = entity.get("vcardArray") or entity.get("vCardArray")
            if not vcard or not isinstance(vcard, list):
                continue
            for prop in _iter_vcard_properties(vcard):
                if not isinstance(prop, list) or len(prop) < 4:
                    continue
                if prop[0] == "adr":
                    # adr: [pobox, ext, street, locality, region, postal, country]
                    values = [v for v in prop[3:] if isinstance(v, str) and v.strip()]
                    if values:
                        return values[-1].strip()
        return None
    return str(data).strip() or None


def _extract_registrar_url(data: Any) -> Optional[str]:
    """
    Извлекает URL сайта регистратора из ответа RDAP/WHOIS.
    RDAP: entities с role registrar -> vcardArray -> url.
    """
    if data is None:
        return None
    if isinstance(data, dict):
        # Прямые ключи
        direct = _get_nested(
            data,
            "registrarUrl",
            "registrar_url",
            "registrarWebsite",
            "registrar_website",
        )
        if direct is not None:
            s = _str_or_first(direct)
            if s and s != "—" and s.startswith(("http://", "https://")):
                return s.strip()
            if s and s != "—":
                # Добавить https если нет протокола
                s = s.strip()
                if s and " " not in s:
                    return f"https://{s}" if "://" not in s else s
        # RDAP entities
        entities = data.get("entities") or data.get("registrar") or []
        if isinstance(entities, dict):
            entities = [entities]
        for entity in entities:
            if not isinstance(entity, dict):
                continue
            roles = entity.get("roles") or entity.get("role") or []
            if isinstance(roles, str):
                roles = [roles]
            if "registrar" not in roles:
                continue
            vcard = entity.get("vcardArray") or entity.get("vCardArray")
            if not vcard or not isinstance(vcard, list):
                continue
            for prop in _iter_vcard_properties(vcard):
                if not isinstance(prop, list) or len(prop) < 4:
                    continue
                if prop[0] == "url":
                    url = prop[3] if len(prop) > 3 else None
                    if url and isinstance(url, str) and url.strip():
                        u = url.strip()
                        if u.startswith(("http://", "https://")):
                            return u
                        return f"https://{u}"
        return None
    return None


def _iter_vcard_properties(vcard: Any) -> list[Any]:
    """Нормализует RDAP vcardArray к списку свойств."""
    if not isinstance(vcard, list):
        return []
    if len(vcard) >= 2 and isinstance(vcard[1], list):
        inner = vcard[1]
        if not inner or isinstance(inner[0], list):
            return inner
    return vcard


def _clean_lookup_value(val: Any) -> Optional[str]:
    """Возвращает осмысленное строковое значение или None."""
    s = _str_or_first(val)
    return s if s and s != "—" else None


def _extract_registrar_name(data: Any) -> Optional[str]:
    """Извлекает имя регистратора из WhoisJSON/RDAP/ViewDNS-подобных ответов."""
    if data is None:
        return None
    if isinstance(data, dict):
        direct = _get_nested(
            data,
            "registrarName",
            "registrar.name",
            "registrar_name",
            "registrar.organization",
            "registrar.org",
            "registrar",
        )
        if isinstance(direct, dict):
            direct = _get_nested(direct, "name", "organization", "org")
        s = _clean_lookup_value(direct)
        if s:
            return s

        entities = data.get("entities") or []
        if isinstance(entities, dict):
            entities = [entities]
        for entity in entities:
            if not isinstance(entity, dict):
                continue
            roles = entity.get("roles") or entity.get("role") or []
            if isinstance(roles, str):
                roles = [roles]
            if "registrar" not in roles:
                continue
            vcard = entity.get("vcardArray") or entity.get("vCardArray")
            if not vcard or not isinstance(vcard, list):
                continue
            for prop in _iter_vcard_properties(vcard):
                if not isinstance(prop, list) or len(prop) < 4:
                    continue
                if prop[0] in ("fn", "org"):
                    value = prop[3] if len(prop) > 3 else None
                    s = _clean_lookup_value(value)
                    if s:
                        return s
        return None
    return _clean_lookup_value(data)


def _extract_rdap_event_date(data: Any, *actions: str) -> Optional[str]:
    """Извлекает дату из RDAP events по eventAction."""
    if not isinstance(data, dict):
        return None
    expected = {a.strip().lower() for a in actions if a.strip()}
    for event in data.get("events") or []:
        if not isinstance(event, dict):
            continue
        action = str(event.get("eventAction") or event.get("action") or "").strip().lower()
        if action not in expected:
            continue
        value = event.get("eventDate") or event.get("date")
        formatted = _format_date(value)
        if formatted != "—":
            return formatted
    return None


def _has_domain_lookup_data(result: dict[str, Any]) -> bool:
    """Проверяет, содержит ли нормализованный ответ полезные поля."""
    return any(
        result.get(key)
        for key in ("created", "expires", "registrar_name", "registrar_country", "registrar_url")
    )


def _normalize_domain_lookup(
    domain: str,
    data: dict[str, Any],
    *,
    source_name: str,
    source_ref: str,
    use_rdap_events: bool = False,
) -> Optional[dict[str, Any]]:
    """Сводит ответы разных сервисов к общей структуре."""
    parsed_domain = _clean_lookup_value(
        _get_nested(
            data,
            "domainName",
            "domain_name",
            "domain",
            "ldhName",
            "unicodeName",
            "name",
        )
    ) or domain
    parsed_domain = parsed_domain.lower()

    created = _clean_lookup_value(
        _get_nested(
            data,
            "created",
            "createdDate",
            "creationDate",
            "created_date",
            "creation_date",
            "standardCreatedDate",
        )
    )
    expires = _clean_lookup_value(
        _get_nested(
            data,
            "expires",
            "expiresDate",
            "expirationDate",
            "expiryDate",
            "expiration_date",
            "expires_date",
            "standardExpiresDate",
        )
    )
    if use_rdap_events:
        created = created or _extract_rdap_event_date(
            data, "registration", "registered", "creation", "created"
        )
        expires = expires or _extract_rdap_event_date(
            data, "expiration", "expiry", "expires", "renewal"
        )

    result = {
        "domain": parsed_domain,
        "created": created,
        "expires": expires,
        "registrar_name": _extract_registrar_name(data),
        "registrar_country": _extract_registrar_country(data),
        "registrar_url": _extract_registrar_url(data),
        "source_name": source_name,
        "source_ref": source_ref,
    }
    return result if _has_domain_lookup_data(result) else None


def _fetch_domain_lookup_whoisjson(domain: str) -> Optional[dict[str, Any]]:
    """Основной lookup через WhoisJSON."""
    if not WHOISJSON_KEY:
        return None
    try:
        r = requests.get(
            WHOISJSON_URL,
            params={"domain": domain},
            headers={
                "Authorization": f"TOKEN={WHOISJSON_KEY}",
                "Accept": "application/json",
            },
            timeout=10,
        )
        r.raise_for_status()
        data = r.json()
        if not isinstance(data, dict):
            return None
        return _normalize_domain_lookup(
            domain,
            data,
            source_name="WhoisJSON",
            source_ref="https://whoisjson.com/",
        )
    except requests.RequestException as e:
        logger.warning("WhoisJSON lookup failed for %s: %s", domain, e)
    except (ValueError, KeyError, TypeError) as e:
        logger.warning("WhoisJSON parse failed for %s: %s", domain, e)
    return None


def _fetch_domain_lookup_rdap(domain: str) -> Optional[dict[str, Any]]:
    """Фолбэк lookup через RDAP.org."""
    try:
        r = requests.get(
            f"{RDAP_DOMAIN_URL}/{domain}",
            headers={"Accept": "application/rdap+json, application/json"},
            timeout=10,
        )
        r.raise_for_status()
        data = r.json()
        if not isinstance(data, dict):
            return None
        return _normalize_domain_lookup(
            domain,
            data,
            source_name="RDAP.org",
            source_ref="https://rdap.org/",
            use_rdap_events=True,
        )
    except requests.RequestException as e:
        logger.warning("RDAP lookup failed for %s: %s", domain, e)
    except (ValueError, KeyError, TypeError) as e:
        logger.warning("RDAP parse failed for %s: %s", domain, e)
    return None


def _extract_viewdns_value(text: str, *patterns: str) -> Optional[str]:
    """Ищет первое совпадение по regex в HTML ViewDNS."""
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE | re.DOTALL)
        if not match:
            continue
        value = html.unescape(match.group(1)).replace("\\/", "/").strip()
        if value and value.lower() != "null":
            return value
    return None


def _fetch_domain_lookup_viewdns(domain: str) -> Optional[dict[str, Any]]:
    """Резервный lookup через публичную HTML-страницу ViewDNS."""
    try:
        r = requests.get(
            VIEWDNS_WHOIS_URL,
            params={"domain": domain},
            headers={
                "Accept": "text/html,application/xhtml+xml",
                "User-Agent": "Mozilla/5.0",
            },
            timeout=12,
        )
        r.raise_for_status()
        text = r.text
    except requests.RequestException as e:
        logger.warning("ViewDNS lookup failed for %s: %s", domain, e)
        return None

    payload = {
        "domainName": _extract_viewdns_value(text, r'"domainName"\s*:\s*"([^"]+)"'),
        "registrarName": _extract_viewdns_value(
            text,
            r'"registrarName"\s*:\s*"([^"]*)"',
            r"Registrar:\s*([^\\\"<]+)",
        ),
        "createdDate": _extract_viewdns_value(
            text,
            r'"standardCreatedDate"\s*:\s*"([^"]+)"',
            r'"createdDate"\s*:\s*"([^"]+)"',
            r"Creation Date:\s*([^\\\"<]+)",
        ),
        "expiresDate": _extract_viewdns_value(
            text,
            r'"standardExpiresDate"\s*:\s*"([^"]+)"',
            r'"expiresDate"\s*:\s*"([^"]+)"',
            r"Expiration Date:\s*([^\\\"<]+)",
            r"Registrar Registration Expiration Date:\s*([^\\\"<]+)",
        ),
        "registrarUrl": _extract_viewdns_value(
            text,
            r"Registrar URL:\s*([^\\\"<]+)",
        ),
    }
    return _normalize_domain_lookup(
        domain,
        payload,
        source_name="ViewDNS.info",
        source_ref="https://viewdns.info/",
    )


def _get_domain_lookup(domain: str) -> Optional[dict[str, Any]]:
    """Оптимизированная цепочка: WhoisJSON -> RDAP.org -> ViewDNS."""
    for fetcher in (
        _fetch_domain_lookup_whoisjson,
        _fetch_domain_lookup_rdap,
        _fetch_domain_lookup_viewdns,
    ):
        result = fetcher(domain)
        if result is not None:
            return result
    return None


def _format_domain_lookup_text(result: dict[str, Any]) -> str:
    """Форматирует нормализованный доменный ответ в HTML."""
    reg_country_line = (
        f"\n{E_COUNTRY} <b>Страна регистратора:</b> {_h(result.get('registrar_country'))}"
        if result.get("registrar_country")
        else ""
    )
    reg_url_line = (
        f'\n{E_LINK} <b>Сайт:</b> <a href="{_h(result.get("registrar_url"))}">{_h(result.get("registrar_url"))}</a>'
        if result.get("registrar_url")
        else ""
    )
    return (
        f"{E_DOMAIN} <b>Информация по домену</b>: <code>{_h(result.get('domain'))}</code>\n"
        f"{DIV}\n"
        f"{E_DATE} <b>Создан:</b> {_h(result.get('created'))}\n"
        f"{E_DATE} <b>Истекает:</b> {_h(result.get('expires'))}\n"
        f"{E_ORG} <b>Регистратор:</b> {_h(result.get('registrar_name'))}"
        f"{reg_country_line}{reg_url_line}"
    )


def get_domain_info(domain: str) -> str:
    """
    Получает информацию о регистрации домена через каскад бесплатных источников:
    WhoisJSON -> RDAP.org -> ViewDNS.
    """
    domain = domain.strip().lower()
    if not _validate_domain(domain):
        return "Неверный формат доменного имени. Пример: example.com"

    result = _get_domain_lookup(domain)
    if result is not None:
        return _format_domain_lookup_text(result)

    return (
        "Не удалось получить WHOIS по этому домену. "
        "Проверьте имя, WHOISJSON_KEY или попробуйте позже."
    )


E_DNS = "🔗"


def _query_dns(domain: str, rtype: str, resolver: Any) -> list[str]:
    """Запрос DNS-записи. Возвращает список значений или пустой список."""
    try:
        answer = resolver.resolve(domain, rtype, lifetime=10)
        result = []
        for rdata in answer:
            if rtype == "A":
                result.append(str(rdata.address))
            elif rtype == "AAAA":
                result.append(str(rdata.address))
            elif rtype == "MX":
                pref = getattr(rdata, "preference", "")
                ex = str(getattr(rdata, "exchange", rdata)).rstrip(".")
                result.append(f"{pref} {ex}" if pref != "" else ex)
            elif rtype in ("NS", "CNAME", "PTR"):
                result.append(str(rdata).rstrip("."))
            elif rtype == "TXT":
                chunks = rdata.strings
                txt = b"".join(c for c in chunks if isinstance(c, bytes)).decode("utf-8", errors="replace") if chunks and isinstance(chunks[0], bytes) else "".join(str(c) for c in chunks)
                result.append(txt)
            else:
                result.append(str(rdata))
        return result
    except Exception:
        return []


def get_dns_info(domain: str) -> str:
    """
    Возвращает DNS-записи домена: A, AAAA, MX, NS, TXT, CNAME.
    Требуется dnspython.
    """
    domain = domain.strip().lower()
    if not _validate_domain(domain):
        return "Неверный формат домена. Пример: example.com"

    if not DNS_AVAILABLE:
        logger.warning("dnspython не установлен")
        return "DNS lookup недоступен. Установите: pip install dnspython"

    resolver = dns.resolver.Resolver()
    resolver.timeout = 8
    resolver.lifetime = 10

    sections = []
    sections.append(f"{E_DNS} <b>DNS записи</b>: <code>{_h(domain)}</code>")
    sections.append(DIV)

    # A
    a_records = _query_dns(domain, "A", resolver)
    if a_records:
        sections.append(f"<b>A</b> (IPv4):\n  " + ", ".join(f"<code>{_h(a)}</code>" for a in a_records))

    # AAAA
    aaaa_records = _query_dns(domain, "AAAA", resolver)
    if aaaa_records:
        sections.append(f"<b>AAAA</b> (IPv6):\n  " + ", ".join(f"<code>{_h(a)}</code>" for a in aaaa_records))

    # MX
    mx_records = _query_dns(domain, "MX", resolver)
    if mx_records:
        sections.append(f"<b>MX</b> (почта):\n  " + "\n  ".join(_h(m) for m in mx_records))

    # NS
    ns_records = _query_dns(domain, "NS", resolver)
    if ns_records:
        sections.append(f"<b>NS</b> (DNS-серверы):\n  " + ", ".join(_h(n) for n in ns_records))

    # CNAME (если есть)
    cname_records = _query_dns(domain, "CNAME", resolver)
    if cname_records:
        sections.append(f"<b>CNAME</b>:\n  " + ", ".join(_h(c) for c in cname_records))

    # TXT (первые 3 для краткости)
    txt_records = _query_dns(domain, "TXT", resolver)
    if txt_records:
        txt_preview = txt_records[:3] if len(txt_records) > 3 else txt_records
        parts = []
        for t in txt_preview:
            s = _h(t)
            parts.append(s[:200] + ("…" if len(s) > 200 else ""))
        txt_str = "\n  ".join(parts)
        if len(txt_records) > 3:
            txt_str += f"\n  … и ещё {len(txt_records) - 3}"
        sections.append(f"<b>TXT</b>:\n  {txt_str}")

    if len(sections) <= 2:
        return "\n".join(sections) + "\n\nЗаписей не найдено или домен не резолвится."

    return "\n".join(sections)


def get_dns_report_file(domain: str) -> tuple[bytes, str]:
    """
    Возвращает DNS-отчёт в виде файла.
    Возвращает (содержимое_байты, имя_файла) или raises.
    """
    domain = domain.strip().lower()
    if not _validate_domain(domain):
        raise ValueError("Неверный формат домена")
    if not DNS_AVAILABLE:
        raise RuntimeError("DNS lookup недоступен. Установите: pip install dnspython")

    resolver = dns.resolver.Resolver()
    resolver.timeout = 8
    resolver.lifetime = 10

    lines = [
        f"DNS ОТЧЁТ: {domain}",
        "=" * 50,
        "",
    ]

    def add_section(title: str, records: list[str]) -> None:
        if records:
            lines.append(f"{title}:")
            for r in records:
                lines.append(f"  {r}")
            lines.append("")

    a_records = _query_dns(domain, "A", resolver)
    add_section("A (IPv4)", a_records)

    aaaa_records = _query_dns(domain, "AAAA", resolver)
    add_section("AAAA (IPv6)", aaaa_records)

    mx_records = _query_dns(domain, "MX", resolver)
    add_section("MX (почта)", mx_records)

    ns_records = _query_dns(domain, "NS", resolver)
    add_section("NS (DNS-серверы)", ns_records)

    cname_records = _query_dns(domain, "CNAME", resolver)
    add_section("CNAME", cname_records)

    txt_records = _query_dns(domain, "TXT", resolver)
    add_section("TXT", txt_records)

    content = "\n".join(lines).encode("utf-8")
    safe_domain = domain.replace(".", "_")
    filename = f"dns_report_{safe_domain}.txt"
    return content, filename


def _format_country(country_val: Any) -> str:
    """Извлекает название страны и регион из ответа API."""
    if country_val is None:
        return "—"
    if isinstance(country_val, dict):
        name = country_val.get("Name") or country_val.get("name") or ""
        region = country_val.get("Cont") or country_val.get("cont") or country_val.get("region") or ""
        if name and region:
            return f"{name} ({region})"
        return str(name).strip() or str(region).strip() or "—"
    return str(country_val).strip() or "—"


def _binlist_json_to_internal(raw: dict[str, Any]) -> dict[str, Any]:
    """Приводит ответ lookup.binlist.net к полям, ожидаемым _format_bin_info."""
    bank = raw.get("bank") if isinstance(raw.get("bank"), dict) else {}
    country = raw.get("country") if isinstance(raw.get("country"), dict) else {}
    scheme = (raw.get("scheme") or "").strip()
    scheme = scheme.title() if scheme else "—"
    ctype = (raw.get("type") or "").strip()
    ctype = ctype.title() if ctype else "—"
    prepaid = raw.get("prepaid")
    if prepaid is True:
        tier = "Prepaid"
    elif prepaid is False:
        tier = "—"
    else:
        tier = "—"
    issuer = (bank.get("name") or "").strip() or "—"
    country_val = (country.get("name") or country.get("alpha2") or "").strip() or "—"
    return {
        "scheme": scheme,
        "type": ctype,
        "Tier": tier,
        "Bank": issuer,
        "Country": country_val,
    }


def _fetch_bin_from_binlist(bin_str: str) -> Optional[dict[str, Any]]:
    """Резервный BIN через публичный API binlist (без ключа)."""
    digits = "".join(c for c in bin_str if c.isdigit())
    if len(digits) < 6:
        return None
    query = digits[:8] if len(digits) >= 8 else digits[:6]
    try:
        r = requests.get(
            f"{BINLIST_LOOKUP_URL}/{query}",
            headers={"Accept": "application/json", "Accept-Version": "3"},
            timeout=15,
        )
        if r.status_code == 404:
            return None
        if r.status_code == 429:
            logger.warning("binlist BIN rate limit (429)")
            return None
        r.raise_for_status()
        raw = r.json()
        if not isinstance(raw, dict):
            return None
        return _binlist_json_to_internal(raw)
    except requests.RequestException as e:
        logger.warning("binlist BIN request failed: %s", e)
        return None
    except (ValueError, TypeError) as e:
        logger.warning("binlist BIN parse error: %s", e)
        return None


def _resolve_bin_data(bin_str: str) -> tuple[Optional[dict[str, Any]], str]:
    """
    Сначала HandyAPI (если есть ключ), иначе сразу binlist.
    Возвращает (данные для _format_bin_info, метка источника для справки).
    """
    if HANDYAPI_KEY:
        try:
            r = requests.get(
                f"{HANDYAPI_BIN_URL}/{bin_str}",
                headers={"x-api-key": HANDYAPI_KEY},
                timeout=15,
            )
            if r.status_code == 200:
                data = r.json()
                if isinstance(data, dict):
                    return data, "https://www.handyapi.com/"
            if r.status_code == 401:
                logger.warning("handyapi BIN unauthorized, пробуем binlist")
            elif r.status_code == 404:
                pass
            else:
                r.raise_for_status()
        except requests.RequestException as e:
            logger.warning("handyapi BIN request failed: %s", e)
        except (ValueError, TypeError) as e:
            logger.warning("handyapi BIN parse error: %s", e)

    bl = _fetch_bin_from_binlist(bin_str)
    if bl is not None:
        return bl, "https://binlist.net/"
    return None, ""


def _format_bin_info(data: dict[str, Any], bin_str: str) -> str:
    """Форматирует ответ HandyAPI BIN в HTML."""
    scheme = _get_nested(data, "Scheme", "scheme", "Brand", "brand") or "—"
    card_type = _get_nested(data, "Type", "type", "CardType", "cardType") or "—"
    tier = _get_nested(data, "Tier", "tier") or "—"
    issuer = _get_nested(data, "Bank", "bank", "Issuer", "issuer") or "—"
    country_raw = _get_nested(data, "Country", "country", "countryName", "country_name")
    country = _format_country(country_raw)

    lines = [
        f"{E_BIN} <b>Информация по BIN</b>: <code>{_h(bin_str)}</code>",
        DIV,
        f"{E_SCHEME} <b>Схема:</b> {_h(scheme)}",
        f"{E_TYPE} <b>Тип карты:</b> {_h(card_type)}",
        f"{E_TYPE} <b>Уровень:</b> {_h(tier)}",
        f"{E_ORG} <b>Банк-эмитент:</b> {_h(issuer)}",
        f"{E_COUNTRY} <b>Страна:</b> {_h(country)}",
    ]
    return "\n".join(lines)


def get_bin_info(bin_str: str) -> str:
    """
    Запрашивает информацию по BIN: сначала HandyAPI (если задан HANDYAPI_KEY),
    при отсутствии результата — резерв lookup.binlist.net (без ключа, лимит запросов).
    BIN — первые 6–8 цифр карты.
    """
    bin_str = bin_str.strip()
    if not _validate_bin(bin_str):
        return "Неверный формат BIN. Укажите 6–8 цифр (например: 535316)."

    data, _ = _resolve_bin_data(bin_str)
    if data is None:
        if HANDYAPI_KEY:
            return (
                "BIN не найден в базе или не удалось получить данные "
                "(HandyAPI и резервный lookup.binlist.net)."
            )
        return (
            "Не удалось получить данные по BIN. "
            "Добавьте HANDYAPI_KEY или попробуйте позже (резерв binlist имеет лимит запросов)."
        )

    return _format_bin_info(data, bin_str)


# Маппинг кодов стран (ISO 3166-1 alpha-2) на полные названия
COUNTRY_NAMES = {
    "RU": "Россия",
    "US": "США",
    "BY": "Беларусь",
    "KZ": "Казахстан",
    "UA": "Украина",
    "DE": "Германия",
    "GB": "Великобритания",
    "FR": "Франция",
    "IT": "Италия",
    "ES": "Испания",
    "PL": "Польша",
    "TR": "Турция",
    "CN": "Китай",
    "IN": "Индия",
    "BR": "Бразилия",
    "JP": "Япония",
    "KR": "Республика Корея",
    "UZ": "Узбекистан",
    "AZ": "Азербайджан",
    "AM": "Армения",
    "GE": "Грузия",
    "MD": "Молдова",
    "LT": "Литва",
    "LV": "Латвия",
    "EE": "Эстония",
    "FI": "Финляндия",
    "SE": "Швеция",
    "NL": "Нидерланды",
    "BE": "Бельгия",
    "AT": "Австрия",
    "CH": "Швейцария",
    "CA": "Канада",
    "AU": "Австралия",
    "IL": "Израиль",
    "AE": "ОАЭ",
    "SA": "Саудовская Аравия",
    "EG": "Египет",
    "TH": "Таиланд",
    "VN": "Вьетнам",
    "ID": "Индонезия",
    "MY": "Малайзия",
    "SG": "Сингапур",
    "PH": "Филиппины",
    "PK": "Пакистан",
    "BD": "Бангладеш",
    "MX": "Мексика",
    "AR": "Аргентина",
    "CL": "Чили",
    "CO": "Колумбия",
    "PE": "Перу",
    "ZA": "ЮАР",
    "NG": "Нигерия",
    "KE": "Кения",
    "GR": "Греция",
    "PT": "Португалия",
    "CZ": "Чехия",
    "RO": "Румыния",
    "HU": "Венгрия",
    "BG": "Болгария",
    "RS": "Сербия",
    "HR": "Хорватия",
    "SK": "Словакия",
    "SI": "Словения",
    "IE": "Ирландия",
    "NZ": "Новая Зеландия",
}

LINE_TYPE_RU = {
    "mobile": "Мобильный",
    "landline": "Стационарный",
    "toll_free": "Бесплатный (8-800)",
    "premium_rate": "Платный",
    "special_services": "Спецслужбы",
    "paging": "Пейджинг",
    "satellite": "Спутниковый",
    "voip": "VoIP",
}


def _format_phone_country(country_name: Optional[str], country_code: Optional[str]) -> str:
    """Возвращает полное название страны. Избегает обрезанных значений вроде (Republic of)."""
    name = (country_name or "").strip()
    code = (country_code or "").strip().upper()
    # Убираем неполные фрагменты в скобках: "(Republic of)", "(Russian Federation)" и т.п.
    if "(" in name and ")" in name:
        name = re.sub(r"\s*\([^)]*\)\s*", " ", name).strip()
    # Неполные или бессмысленные названия от API
    incomplete = (
        not name
        or len(name) <= 3
        or name.lower() in ("republic of", "(republic of)", "republic of")
    )
    if incomplete and code:
        return COUNTRY_NAMES.get(code, code)
    if name and not incomplete:
        return name
    if code:
        return COUNTRY_NAMES.get(code, code)
    return "—"


def _format_phone_info(data: dict[str, Any], phone: str) -> str:
    """Форматирует ответ Numverify в HTML."""
    carrier = data.get("carrier") or "—"
    line_type_en = data.get("line_type") or ""
    line_type = LINE_TYPE_RU.get(line_type_en.lower(), line_type_en) if line_type_en else "—"
    country = _format_phone_country(
        data.get("country_name"),
        data.get("country_code"),
    )
    location = data.get("location") or "—"
    intl_format = data.get("international_format") or phone
    valid = data.get("valid", True)

    lines = [
        f"{E_PHONE} <b>Информация по номеру</b>: <code>{_h(intl_format)}</code>",
        DIV,
        f"{E_ORG} <b>Оператор:</b> {_h(carrier)}",
        f"{E_TYPE} <b>Тип линии:</b> {_h(line_type)}",
        f"{E_COUNTRY} <b>Страна:</b> {_h(country)}",
        f"{E_CITY} <b>Регион/город:</b> {_h(location)}",
    ]
    if not valid:
        lines.insert(2, "⚠️ <i>Номер не найден или недействителен</i>")

    return "\n".join(lines)


def get_phone_info(phone: str) -> str:
    """
    Запрашивает информацию об операторе по номеру через Numverify.
    Требуется NUMVERIFY_KEY (1000 бесплатных запросов/мес на numverify.com).
    """
    phone = phone.strip()
    if not _validate_phone(phone):
        return "Неверный формат номера. Укажите 10–15 цифр с кодом страны (например: +79161234567)."

    normalized = _normalize_phone(phone)
    if not normalized:
        return "Не удалось распознать номер."

    if not NUMVERIFY_KEY:
        logger.warning("NUMVERIFY_KEY не задан")
        return "API номера не настроен. Введите API ключ в ip_service.py (NUMVERIFY_KEY)."

    # Numverify ожидает номер без + (только цифры)
    number_param = re.sub(r"\D", "", normalized)

    try:
        r = requests.get(
            NUMVERIFY_URL,
            params={"access_key": NUMVERIFY_KEY, "number": number_param},
            timeout=15,
        )
        r.raise_for_status()
        data = r.json()
    except requests.RequestException as e:
        logger.warning("numverify request failed: %s", e)
        if hasattr(e, "response") and e.response is not None:
            status = getattr(e.response, "status_code", None)
            if status == 401:
                return "Неверный API-ключ Numverify. Проверьте NUMVERIFY_KEY."
        return "Не удалось получить данные по номеру. Попробуйте позже."
    except (ValueError, KeyError) as e:
        logger.warning("numverify parse error: %s", e)
        return "Ошибка при разборе ответа API. Попробуйте позже."

    if data.get("success") is False:
        err = data.get("error", {})
        info = err.get("info", "Неизвестная ошибка API.")
        if err.get("type") == "invalid_access_key":
            return "Неверный API-ключ Numverify. Проверьте NUMVERIFY_KEY."
        if err.get("type") == "usage_limit_reached":
            return "Достигнут лимит запросов Numverify. Попробуйте в следующем месяце."
        return f"Ошибка API: {info}"

    return _format_phone_info(data, normalized)


# ============ Криптокошелёк (ETH, BTC, TRON) ============

def _wallet_label(addr: str, chain: str) -> str:
    """Метка для адреса из WALLET_LABELS."""
    key = addr.lower().strip()
    return WALLET_LABELS.get(key, "—")


def _fetch_eth_wallet(address: str) -> Optional[dict[str, Any]]:
    """ETH через Etherscan. Требуется ETHERSCAN_KEY."""
    if not ETHERSCAN_KEY:
        return None
    addr = address.strip()
    try:
        # balance
        r = requests.get(
            ETHERSCAN_API_URL,
            params={
                "module": "account",
                "action": "balance",
                "address": addr,
                "tag": "latest",
                "apikey": ETHERSCAN_KEY,
            },
            timeout=15,
        )
        r.raise_for_status()
        data = r.json()
        if data.get("status") != "1" and data.get("message") != "OK":
            return None
        balance_wei = int(data.get("result", 0))
        balance_eth = balance_wei / 1e18

        # tx count (макс 10000 за запрос)
        r2 = requests.get(
            ETHERSCAN_API_URL,
            params={
                "module": "account",
                "action": "txlist",
                "address": addr,
                "startblock": 0,
                "endblock": 99999999,
                "page": 1,
                "offset": 10000,
                "sort": "asc",
                "apikey": ETHERSCAN_KEY,
            },
            timeout=20,
        )
        tx_count = 0
        if r2.status_code == 200:
            data2 = r2.json()
            if data2.get("status") == "1" and isinstance(data2.get("result"), list):
                tx_count = len(data2["result"])
                if tx_count >= 10000:
                    tx_count = 10000  # Etherscan макс, реальное может быть больше

        return {
            "chain": "eth",
            "balance": balance_eth,
            "balance_raw": balance_wei,
            "tx_count": tx_count,
            "symbol": "ETH",
            "label": _wallet_label(addr, "eth"),
        }
    except Exception as e:
        logger.warning("eth wallet fetch failed: %s", e)
        return None


def _fetch_btc_wallet(address: str) -> Optional[dict[str, Any]]:
    """BTC через BlockCypher (без ключа)."""
    addr = address.strip()
    try:
        r = requests.get(
            f"{BLOCKCYPHER_URL}/btc/main/addrs/{addr}",
            params={"limit": 1},
            timeout=15,
        )
        r.raise_for_status()
        data = r.json()
        balance_satoshi = data.get("balance", 0) or data.get("final_balance", 0)
        tx_count = data.get("n_tx", 0) or data.get("final_n_tx", 0)
        balance_btc = balance_satoshi / 1e8
        return {
            "chain": "btc",
            "balance": balance_btc,
            "balance_raw": balance_satoshi,
            "tx_count": tx_count,
            "symbol": "BTC",
            "label": _wallet_label(addr, "btc"),
        }
    except Exception as e:
        logger.warning("btc wallet fetch failed: %s", e)
        return None


def _fetch_tron_wallet(address: str) -> Optional[dict[str, Any]]:
    """TRON через TronGrid."""
    addr = address.strip()
    try:
        r = requests.get(
            f"{TRONGRID_URL.rstrip('/')}/v1/accounts/{addr}",
            timeout=15,
        )
        r.raise_for_status()
        data = r.json()
        accs = data.get("data") or data.get("accounts") or []
        if not accs:
            return {"chain": "tron", "balance": 0, "tx_count": 0, "symbol": "TRX", "label": "—"}
        acc = accs[0] if isinstance(accs[0], dict) else accs
        balance_sun = int(acc.get("balance", 0))
        balance_trx = balance_sun / 1e6
        # tx count — нужен отдельный запрос
        r2 = requests.get(
            f"{TRONGRID_URL.rstrip('/')}/v1/accounts/{addr}/transactions",
            params={"limit": 1},
            timeout=15,
        )
        tx_count = 0
        if r2.status_code == 200:
            j = r2.json()
            meta = j.get("meta", {})
            tx_count = meta.get("page_size", 0)  # нет total, используем макс через pagination
        # TronGrid не возвращает total в meta — делаем ещё запрос с only_confirmed
        r3 = requests.get(
            f"{TRONGRID_URL.rstrip('/')}/v1/accounts/{addr}/transactions",
            params={"limit": 200, "only_confirmed": "true"},
            timeout=15,
        )
        if r3.status_code == 200:
            j3 = r3.json()
            txs = j3.get("data") or []
            tx_count = len(txs)  # хотя бы из первой страницы
        return {
            "chain": "tron",
            "balance": balance_trx,
            "balance_raw": balance_sun,
            "tx_count": tx_count,
            "symbol": "TRX",
            "label": _wallet_label(addr.lower(), "tron"),
        }
    except Exception as e:
        logger.warning("tron wallet fetch failed: %s", e)
        return None


def _format_wallet_amount(val: float, symbol: str) -> str:
    if val >= 1e9:
        return f"{val:,.2f} {symbol}"
    if val >= 1:
        return f"{val:,.4f} {symbol}"
    if val >= 1e-6:
        return f"{val:.8f} {symbol}"
    return f"{val} {symbol}"


def get_wallet_info(address: str) -> str:
    """
    Мультичейн: ETH, BTC, TRON.
    Баланс, количество транзакций, метки.
    """
    addr = address.strip()
    if not _validate_wallet(addr):
        return "Неверный формат адреса. Поддерживаются: ETH (0x...), BTC (1..., 3..., bc1...), TRON (T...)."

    chains = _detect_wallet_chains(addr)
    if not chains:
        return "Адрес не распознан. Поддерживаются: ETH, BTC, TRON."

    results: list[dict[str, Any]] = []
    if "eth" in chains:
        r = _fetch_eth_wallet(addr)
        if r:
            results.append(r)
        elif ETHERSCAN_KEY:
            results.append({"chain": "eth", "balance": 0, "tx_count": 0, "symbol": "ETH", "label": "—", "error": True})
    if "btc" in chains:
        r = _fetch_btc_wallet(addr)
        if r:
            results.append(r)
    if "tron" in chains:
        r = _fetch_tron_wallet(addr)
        if r:
            results.append(r)

    if not results:
        eth_hint = " Добавьте ETHERSCAN_KEY в ip_service.py для ETH." if "eth" in chains else ""
        return f"Не удалось получить данные по адресу.{eth_hint}"

    chain_names = {"eth": "Ethereum", "btc": "Bitcoin", "tron": "TRON"}
    lines = [
        f"{E_WALLET} <b>Криптокошелёк</b>: <code>{_h(addr)}</code>",
        DIV,
    ]
    for r in results:
        chain = chain_names.get(r["chain"], r["chain"].upper())
        bal = _format_wallet_amount(r.get("balance", 0), r.get("symbol", ""))
        tx_count = r.get("tx_count", 0)
        label = r.get("label", "—")
        lines.append(f"<b>{chain}</b>")
        lines.append(f"  {E_STAR} Баланс: {_h(bal)}")
        lines.append(f"  {E_TYPE} Транзакций: {_h(tx_count)}")
        lines.append(f"  {E_ORG} Метка: {_h(label)}")
        lines.append("")
    if "eth" in chains and not any(r.get("chain") == "eth" for r in results) and not ETHERSCAN_KEY:
        lines.append("<i>ETH: укажите ETHERSCAN_KEY в ip_service.py</i>")

    return "\n".join(lines).strip()


MAX_TX_IN_REPORT = 1000


def _collect_eth_transactions(address: str) -> list[dict[str, Any]]:
    """Собирает транзакции ETH (Etherscan)."""
    if not ETHERSCAN_KEY:
        return []
    txs = []
    page = 1
    offset = 1000
    while len(txs) < MAX_TX_IN_REPORT:
        try:
            r = requests.get(
                ETHERSCAN_API_URL,
                params={
                    "module": "account",
                    "action": "txlist",
                    "address": address,
                    "startblock": 0,
                    "endblock": 99999999,
                    "page": page,
                    "offset": offset,
                    "sort": "desc",
                    "apikey": ETHERSCAN_KEY,
                },
                timeout=25,
            )
            r.raise_for_status()
            data = r.json()
            if data.get("status") != "1":
                break
            batch = data.get("result") or []
            if not batch:
                break
            txs.extend(batch)
            if len(batch) < offset:
                break
            page += 1
        except Exception:
            break
    return txs[:MAX_TX_IN_REPORT]


def _collect_btc_transactions(address: str) -> list[dict[str, Any]]:
    """Собирает транзакции BTC (BlockCypher)."""
    try:
        r = requests.get(
            f"{BLOCKCYPHER_URL}/btc/main/addrs/{address}",
            params={"limit": 500, "unspentOnly": "false"},
            timeout=30,
        )
        r.raise_for_status()
        data = r.json()
        refs = data.get("txrefs") or data.get("unconfirmed_txrefs") or []
        txs = []
        for ref in refs[:MAX_TX_IN_REPORT]:
            txs.append({
                "hash": ref.get("tx_hash"),
                "block": ref.get("block_height"),
                "date": ref.get("confirmed"),
                "value": ref.get("value", 0),
                "fee": ref.get("fees", 0),
                "spent": ref.get("spent", False),  # True = исходящая
            })
        return txs
    except Exception:
        return []


def _collect_tron_transactions(address: str) -> list[dict[str, Any]]:
    """Собирает транзакции TRON (TronGrid)."""
    from datetime import datetime
    txs = []
    fingerprint = ""
    while len(txs) < MAX_TX_IN_REPORT:
        try:
            params: dict[str, str] = {"limit": "200", "only_confirmed": "true"}
            if fingerprint:
                params["fingerprint"] = fingerprint
            r = requests.get(
                f"{TRONGRID_URL.rstrip('/')}/v1/accounts/{address}/transactions",
                params=params,
                timeout=20,
            )
            r.raise_for_status()
            data = r.json()
            batch = data.get("data") or []
            if not batch:
                break
            for b in batch:
                tx_id = b.get("txID") or b.get("tx_hash") or "—"
                block = b.get("blockNumber") or b.get("block") or "—"
                ts = b.get("block_timestamp") or b.get("timestamp") or 0
                ts_str = "—"
                if ts:
                    try:
                        dt = datetime.fromtimestamp(ts / 1000)
                        ts_str = dt.strftime("%Y-%m-%d %H:%M")
                    except Exception:
                        ts_str = str(ts)
                # Извлекаем from/to из raw_data.contract
                from_addr = to_addr = None
                try:
                    raw = b.get("raw_data") or {}
                    contracts = raw.get("contract") or []
                    if contracts:
                        param = contracts[0].get("parameter", {}) if isinstance(contracts[0], dict) else {}
                        val = (param.get("value") or {}) if isinstance(param, dict) else {}
                        from_addr = val.get("owner_address") or val.get("ownerAddress")
                        to_addr = val.get("to_address") or val.get("toAddress")
                except Exception:
                    pass
                txs.append({
                    "hash": tx_id, "block": block, "date": ts_str, "value": 0,
                    "from": from_addr, "to": to_addr,
                })
            fingerprint = data.get("meta", {}).get("fingerprint") or ""
            if not fingerprint or len(batch) < 200:
                break
        except Exception:
            break
    return txs[:MAX_TX_IN_REPORT]


def get_wallet_tx_report_file(address: str, chain: str) -> tuple[bytes, str]:
    """
    Возвращает файл со всеми транзакциями.
    chain: eth, btc, tron.
    """
    addr = address.strip()
    if not _validate_wallet(addr):
        raise ValueError("Неверный формат адреса кошелька")
    if chain not in ("eth", "btc", "tron"):
        raise ValueError(f"Сеть не поддерживается: {chain}")

    if chain == "eth":
        txs = _collect_eth_transactions(addr)
    elif chain == "btc":
        txs = _collect_btc_transactions(addr)
    else:
        txs = _collect_tron_transactions(addr)

    symbols = {"eth": "ETH", "btc": "BTC", "tron": "TRX"}
    divs = {"eth": 1e18, "btc": 1e8, "tron": 1e6}
    sym = symbols.get(chain, "")
    div = divs.get(chain, 1)

    lines = [
        f"Транзакции: {addr}",
        f"Сеть: {chain.upper()}",
        "=" * 60,
        "",
    ]
    addr_lower = addr.lower()
    for i, t in enumerate(txs, 1):
        h = t.get("hash") or t.get("tx_hash") or "—"
        blk = t.get("block") or t.get("block_height") or t.get("blockNumber") or "—"
        dt = t.get("date") or t.get("confirmed") or t.get("timeStamp")
        if dt and isinstance(dt, (int, float)):
            try:
                from datetime import datetime
                dt = datetime.fromtimestamp(int(dt)).strftime("%Y-%m-%d %H:%M")
            except Exception:
                dt = str(dt)
        dt = dt or "—"
        val = t.get("value") or 0
        if isinstance(val, (int, float)) and div:
            val_fmt = f"{val / div:.8f} {sym}"
        else:
            val_fmt = str(val)
        lines.append(f"{i}. {h}")
        lines.append(f"   Block: {blk} | Date: {dt} | Value: {val_fmt}")

        # Адреса from/to — откуда пришли или куда ушли токены (контрагенты)
        from_addr = t.get("from") or t.get("fromAddress")
        to_addr = t.get("to") or t.get("toAddress")
        if from_addr or to_addr:
            from_str = (from_addr or "").strip()
            to_str = (to_addr or "").strip()
            lines.append(f"   From: {from_str or '—'}")
            lines.append(f"   To: {to_str or '—'}")
            # Подсветка контрагента (второй участник) для ETH
            if from_str and to_str and chain == "eth" and addr:
                if from_str.lower() == addr_lower:
                    lines.append(f"   Контрагент (получатель): {to_str}")
                elif to_str.lower() == addr_lower:
                    lines.append(f"   Контрагент (отправитель): {from_str}")
        # BTC: txrefs не содержат from/to, но spent показывает направление
        elif "spent" in t:
            tx_hash = t.get("hash") or t.get("tx_hash")
            direction = "исходящая" if t.get("spent") else "входящая"
            link = f" (детали: blockcypher.com/btc/tx/{tx_hash})" if tx_hash else ""
            lines.append(f"   Направление: {direction}{link}")
        lines.append("")

    tail = f"\n(Показано {len(txs)} транзакций)" if txs else "\nТранзакций не найдено."
    content = ("\n".join(lines) + tail).encode("utf-8")
    safe_addr = "".join(c if c.isalnum() else "_" for c in addr[:30])
    filename = f"tx_{chain}_{safe_addr or 'addr'}.txt"
    return content, filename


# ============ Справка для уголовного дела (Word) ============

def _spravka_pl(val: str, default: str = "_______________") -> str:
    """Подставное значение или placeholder."""
    return (val or "").strip() or default


def _spravka_unit_genitive(unit: str) -> str:
    """Подразделение в родительном падеже: Дзержинский РОСК → Дзержинского РОСК."""
    unit = (unit or "").strip()
    if not unit:
        return unit
    parts = unit.split(None, 1)
    word = parts[0]
    rest = (" " + parts[1]) if len(parts) > 1 else ""
    if word.endswith("ский"):
        word = word[:-4] + "ского"
    elif word.endswith("ий"):
        word = word[:-2] + "ого"
    return word + rest


def _spravka_unit_signature_lines(unit: str) -> tuple[str, str, str, str]:
    """
    Развёртка подразделения для подписи (4 строки, РОСК/ГОСК).
    Возвращает (прилагательное в род. п., "районного/городского отдела", "Следственного комитета", "Республики Беларусь").
    """
    unit = (unit or "").strip()
    if not unit:
        return "_______________", "районного отдела", "Следственного комитета", "Республики Беларусь"
    parts = unit.split(None, 1)
    word = parts[0]
    is_gosk = len(parts) > 1 and "ГОСК" in (parts[1] or "")
    if word.endswith("ский"):
        first = word[:-4] + "ского"
    elif word.endswith("ий"):
        first = word[:-2] + "ого"
    else:
        first = word
    second = "городского отдела" if is_gosk else "районного отдела"
    return first, second, "Следственного комитета", "Республики Беларусь"


def _spravka_signature_unit_paragraphs(unit: str, position: str) -> list[str]:
    """
    Список строк подписи по подразделению и должности.
    СУ УСК и ОЦРПС УСК — 5 строк; РОСК/ГОСК — 4 строки.
    """
    unit = (unit or "").strip()
    pos = (position or "").strip()
    if unit == "СУ УСК":
        return [
            f"{pos} отдела",
            "следственного управления",
            "управления Следственного комитета",
            "Республики Беларусь",
            "по Минской области",
        ]
    if unit == "ОЦРПС УСК":
        return [
            f"{pos} отдела",
            "цифрового развития",
            "предварительного следствия",
            "управления Следственного комитета",
            "Республики Беларусь",
            "по Минской области",
        ]
    first, second, third, fourth = _spravka_unit_signature_lines(unit)
    return [f"{pos} {first}", second, third, fourth]


def get_spravka_word(
    lookup_type: str,
    value: str,
    case_num: str | None = None,
    position: str | None = None,
    unit: str | None = None,
    rank: str | None = None,
    signature_name: str | None = None,
) -> tuple[bytes, str]:
    """
    Формирует справку в формате Word (.docx) по образцу для уголовного дела.
    Возвращает (содержимое_файла, имя_файла).
    lookup_type: ip, domain, dns, bin, phone, wallet
    case_num, position, unit, rank, signature_name — подставляются в подпись при передаче.
    """
    from datetime import datetime
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("Установите python-docx: pip install python-docx")

    doc = Document()
    # Поля листа: сверху и снизу 2 см, слева 3 см, справа 1 см
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1)

    # Стиль основного текста: 15 pt, Times New Roman, одинарный интервал, отступ первой строки 1,25 см, без интервала после абзаца
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(15)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0  # одинарный
    style.paragraph_format.first_line_indent = Cm(1.25)

    case_num = (case_num or "").strip() or _spravka_pl(SPRAVKA_CASE, "_______________")
    today = datetime.now().strftime("%d.%m.%Y")
    # Подпись: из параметров или константы
    pos_val = (position or "").strip() or _spravka_pl(SPRAVKA_POSITION, "_______________")
    unit_val = (unit or "").strip()
    rank_val = (rank or "").strip() or _spravka_pl(SPRAVKA_RANK, "_______________")
    name_val = (signature_name or "").strip() or _spravka_pl(SPRAVKA_NAME, "_______________")

    # Заголовок: по центру, 15 pt, Times New Roman, межстрочный 14 pt, без жирного, без интервала после абзацев
    # Слово «СПРАВКА» — отдельной строкой заглавными
    p1 = doc.add_paragraph("СПРАВКА")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.line_spacing = Pt(14)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.runs[0]
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(15)

    p2 = doc.add_paragraph("по результатам проверки идентификаторов с использованием открытых интернет-источников")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = Pt(14)
    p2.paragraph_format.space_after = Pt(0)
    for r in p2.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(15)

    # Пустой абзац-отступ между заголовком и текстом (межстрочный 1,5)
    p_gap = doc.add_paragraph()
    p_gap.paragraph_format.space_after = Pt(0)
    p_gap.paragraph_format.line_spacing = 1.5

    # Вводная часть — зависит от типа
    type_names = {
        "ip": ("IP-адреса", value),
        "domain": ("доменного имени", value),
        "dns": ("домена", value),
        "bin": ("BIN банковской карты", value),
        "phone": ("номера телефона", value),
        "wallet": ("криптовалютного адреса", value),
    }
    id_type, id_val = type_names.get(lookup_type, ("идентификатора", value))
    intro = (
        f"В рамках расследования уголовного дела № {case_num.strip()} посредством открытых "
        f"интернет-источников осуществлена проверка {id_type} {id_val}, по результатам которой "
        "получена следующая информация, имеющая значение для расследования уголовного дела."
    )
    p_intro = doc.add_paragraph(intro)
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    source_url = ""
    ref_osint = False  # при True в Справочно — формулировка про OSINT-методы вместо интернет-сервиса
    content_lines: list[str | None] = []

    if lookup_type == "ip":
        data = _fetch_ip_data_raw(value)
        if not data:
            raise ValueError("Не удалось получить данные по IP")
        loc = data.get("location") or {}
        asn_data = data.get("asn") or {}
        country = loc.get("country", "—")
        city = loc.get("city", "—")
        org = asn_data.get("org", "—")
        asn_num = asn_data.get("asn", "")
        provider = f"{org} ({asn_num})" if asn_num else org
        yd = _yes_no
        content_lines = [
            f"Страна: {country}",
            f"Город: {city or '–'}",
            f"Провайдер (ASN): {provider or '–'}",
            None,  # отступ (межстрочный 1,5) между Провайдер и Безопасность
            "Безопасность:",
            f"Дата-центр: {yd(data.get('is_datacenter', False)).lower()}",
            f"VPN: {yd(data.get('is_vpn', False)).lower()}",
            f"Прокси: {yd(data.get('is_proxy', False)).lower()}",
            f"TOR: {yd(data.get('is_tor', False)).lower()}",
        ]
        source_url = API_URL

    elif lookup_type == "domain":
        domain = value.strip().lower()
        result = _get_domain_lookup(domain)
        if result is None:
            raise ValueError("Не удалось получить данные по домену")
        content_lines = [
            f"Создан: {result.get('created') or '–'}",
            f"Истекает: {result.get('expires') or '–'}",
            f"Регистратор: {result.get('registrar_name') or '–'}",
        ]
        source_url = result.get("source_name") or result.get("source_ref") or "—"

    elif lookup_type == "bin":
        data, bin_source = _resolve_bin_data(value.strip())
        if not data:
            raise ValueError("Не удалось получить данные по BIN")
        scheme = _get_nested(data, "Scheme", "scheme") or "—"
        card_type = _get_nested(data, "Type", "type") or "—"
        issuer = _get_nested(data, "Bank", "Issuer", "bank", "issuer") or "—"
        country = _format_country(data.get("Country") or data.get("country"))
        content_lines = [
            f"Схема: {scheme}",
            f"Тип: {card_type}",
            f"Банк-эмитент: {issuer}",
            f"Страна: {country}",
        ]
        source_url = bin_source or "—"

    elif lookup_type == "phone":
        if not NUMVERIFY_KEY:
            raise ValueError("API номера не настроен. Укажите NUMVERIFY_KEY в ip_service.py.")
        number_param = re.sub(r"\D", "", value)
        try:
            r = requests.get(
                NUMVERIFY_URL,
                params={"access_key": NUMVERIFY_KEY, "number": number_param},
                timeout=15,
            )
            r.raise_for_status()
            data = r.json()
        except requests.RequestException as e:
            logger.warning("numverify request failed for spravka: %s", e)
            raise ValueError("Не удалось получить данные по номеру телефона.")
        if data.get("success") is False:
            err = data.get("error", {})
            raise ValueError(f"Ошибка API номера: {err.get('info', 'Неизвестная ошибка API')}")
        carrier = data.get("carrier") or "—"
        line_type_en = data.get("line_type") or ""
        line_type = LINE_TYPE_RU.get(line_type_en.lower(), line_type_en) if line_type_en else "—"
        country = _format_phone_country(data.get("country_name"), data.get("country_code"))
        location = data.get("location") or "—"
        intl_format = data.get("international_format") or value
        content_lines = [
            f"Номер: {intl_format}",
            f"Оператор: {carrier}",
            f"Тип линии: {line_type}",
            f"Страна: {country}",
            f"Регион/город: {location}",
        ]
        source_url = NUMVERIFY_URL

    elif lookup_type == "wallet":
        content_lines = [
            f"Проверяемое значение: {value}",
            "(Данные получены через OSINT-инструмент. Сведения уточните по результатам проверки.)",
        ]
        chains = _detect_wallet_chains(value)
        parts = []
        if "eth" in chains:
            parts.append("Etherscan (etherscan.io)")
        if "btc" in chains:
            parts.append("BlockCypher (blockcypher.com)")
        if "tron" in chains:
            parts.append("TronGrid (trongrid.io)")
        source_url = ", ".join(parts) if parts else "—"

    elif lookup_type == "dns":
        content_lines = [
            f"Проверяемое значение: {value}",
            "(Данные получены через OSINT-инструмент. Сведения уточните по результатам проверки.)",
        ]
        source_url = "DNS (системный резолвер)"

    else:
        content_lines = [
            f"Проверяемое значение: {value}",
            "(Данные получены через OSINT-инструмент. Сведения уточните по результатам проверки.)",
        ]
        source_url = "—"

    # Отступ между вводным абзацем и блоком данных (межстрочный 1,5)
    p_gap2 = doc.add_paragraph()
    p_gap2.paragraph_format.space_after = Pt(0)
    p_gap2.paragraph_format.line_spacing = 1.5

    for line in content_lines:
        if line is None:
            p_gap = doc.add_paragraph()
            p_gap.paragraph_format.space_after = Pt(0)
            p_gap.paragraph_format.line_spacing = 1.5
            continue
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Отступ между блоком данных и «Справочно» (межстрочный 1,5)
    p_gap3 = doc.add_paragraph()
    p_gap3.paragraph_format.space_after = Pt(0)
    p_gap3.paragraph_format.line_spacing = 1.5

    # Справочно (курсив)
    if ref_osint:
        ref_text = "Справочно: информация получена с использованием OSINT-методов поиска информации."
    else:
        ref_text = f"Справочно: информация получена с использованием интернет-сервиса «{source_url}»."
    p_ref = doc.add_paragraph(ref_text)
    p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p_ref.runs:
        r.font.italic = True

    # Подпись: без отступа первой строки (1,25 см); межстрочный 14 pt
    sig_line_spacing = Pt(14)
    sig_no_indent = Pt(0)
    p_empty = doc.add_paragraph()
    p_empty.paragraph_format.first_line_indent = sig_no_indent
    if unit_val and pos_val and rank_val and name_val:
        unit_lines = _spravka_signature_unit_paragraphs(unit_val, pos_val)
        for line in unit_lines:
            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing = sig_line_spacing
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = sig_no_indent
        # Звание уже с «юстиции» из бота; 6 Tab до ФИО
        num_tabs = 6
        tabs_str = "\t" * num_tabs
        rank_line = doc.add_paragraph()
        rank_line.paragraph_format.line_spacing = sig_line_spacing
        rank_line.paragraph_format.space_after = Pt(0)
        rank_line.paragraph_format.first_line_indent = sig_no_indent
        r = rank_line.add_run(f"{rank_val}{tabs_str}{name_val}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(15)
    else:
        p_old = doc.add_paragraph(pos_val)
        p_old.paragraph_format.line_spacing = sig_line_spacing
        p_old.paragraph_format.space_after = Pt(0)
        p_old.paragraph_format.first_line_indent = sig_no_indent
        p_old = doc.add_paragraph(rank_val)
        p_old.paragraph_format.line_spacing = sig_line_spacing
        p_old.paragraph_format.space_after = Pt(0)
        p_old.paragraph_format.first_line_indent = sig_no_indent
        p_old = doc.add_paragraph(name_val)
        p_old.paragraph_format.line_spacing = sig_line_spacing
        p_old.paragraph_format.space_after = Pt(0)
        p_old.paragraph_format.first_line_indent = sig_no_indent
    p_date = doc.add_paragraph(today)
    p_date.paragraph_format.line_spacing = sig_line_spacing
    p_date.paragraph_format.space_after = Pt(0)
    p_date.paragraph_format.first_line_indent = sig_no_indent

    # Сохраняем в bytes
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    content = buffer.read()

    safe_val = "".join(c if c.isalnum() else "_" for c in str(value)[:20])
    filename = f"Справка_{lookup_type}_{safe_val}.docx"
    return content, filename
